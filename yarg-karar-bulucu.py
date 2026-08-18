import io
import re
import html
import json
import time
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import quote

import requests
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt


# ============================================================
# SAYFA AYARLARI
# ============================================================

st.set_page_config(
    page_title="Yargı Kararı Bulucu",
    page_icon="⚖️",
    layout="wide",
)

SOURCES = {
    "Yargıtay": {
        "base": "https://karararama.yargitay.gov.tr",
        "search": "https://karararama.yargitay.gov.tr/aramadetaylist",
        "doc": (
            "https://karararama.yargitay.gov.tr/"
            "getDokuman?id={id}&arananKelime={term}"
        ),
    },
    "UYAP Emsal": {
        "base": "https://emsal.uyap.gov.tr",
        "search": "https://emsal.uyap.gov.tr/aramadetaylist",
        "doc": (
            "https://emsal.uyap.gov.tr/"
            "getDokuman?arananKelime={term}&id={id}"
        ),
    },
}

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0 Safari/537.36"
    ),
    "Accept": "application/json, text/plain, */*",
    "Accept-Language": "tr-TR,tr;q=0.9,en;q=0.7",
}


# ============================================================
# RUNTIME STATE
# ============================================================

@st.cache_resource
def runtime_state():
    return {
        "sessions": {},
        "locks": {
            "Yargıtay": threading.Lock(),
            "UYAP Emsal": threading.Lock(),
        },
        "last_request": {
            "Yargıtay": 0.0,
            "UYAP Emsal": 0.0,
        },
        "schema": {},
    }


def get_session(source):
    state = runtime_state()

    if source not in state["sessions"]:
        cfg = SOURCES[source]

        session = requests.Session()
        session.headers.update({
            **HEADERS,
            "Content-Type": "application/json;charset=UTF-8",
            "Origin": cfg["base"],
            "Referer": cfg["base"] + "/",
        })

        state["sessions"][source] = session

    return state["sessions"][source]


def request_source(source, method, url, **kwargs):
    state = runtime_state()
    lock = state["locks"][source]

    with lock:
        elapsed = time.monotonic() - state["last_request"][source]

        if elapsed < 1.2:
            time.sleep(1.2 - elapsed)

        session = get_session(source)
        last_error = None

        for attempt in range(2):
            try:
                response = session.request(
                    method,
                    url,
                    timeout=(5, 20),
                    **kwargs,
                )

                state["last_request"][source] = time.monotonic()

                if response.status_code == 429 and attempt == 0:
                    time.sleep(6)
                    continue

                response.raise_for_status()
                return response

            except Exception as exc:
                last_error = exc

                if attempt == 0:
                    time.sleep(1.5)

        raise last_error


# ============================================================
# YARDIMCILAR
# ============================================================

def clean(value):
    return re.sub(r"\s+", " ", str(value or "")).strip()


def parse_number(value):
    try:
        return int(
            str(value)
            .replace(".", "")
            .replace(",", "")
            .strip()
        )
    except Exception:
        return None


def pick(data, *keys):
    if not isinstance(data, dict):
        return ""

    lowered = {
        str(k).lower(): v
        for k, v in data.items()
    }

    for key in keys:
        if key in data and data[key] not in (None, ""):
            return data[key]

        value = lowered.get(str(key).lower())

        if value not in (None, ""):
            return value

    return ""


# ============================================================
# JSON KARAR LİSTESİ BULMA
# ============================================================

def looks_like_decision(row):
    if not isinstance(row, dict):
        return False

    keys = " ".join(
        str(k).lower()
        for k in row.keys()
    )

    return any(
        token in keys
        for token in [
            "esas",
            "karar",
            "daire",
            "birim",
            "dokuman",
            "document",
            "tarih",
        ]
    )


def recursive_lists(obj):
    found = []

    if isinstance(obj, list):
        if obj and all(isinstance(x, dict) for x in obj):
            found.append(obj)

        for item in obj:
            found.extend(recursive_lists(item))

    elif isinstance(obj, dict):
        for value in obj.values():
            found.extend(recursive_lists(value))

    return found


def find_rows(obj):
    common_paths = [
        ("data", "data"),
        ("data", "content"),
        ("data", "rows"),
        ("data", "list"),
        ("data", "result"),
        ("rows",),
        ("content",),
        ("result",),
    ]

    for path in common_paths:
        current = obj
        valid = True

        for key in path:
            if not isinstance(current, dict) or key not in current:
                valid = False
                break

            current = current[key]

        if valid and isinstance(current, list):
            if not current or any(
                looks_like_decision(x)
                for x in current
            ):
                return current

    candidates = recursive_lists(obj)

    candidates.sort(
        key=lambda lst: sum(
            1 for row in lst
            if looks_like_decision(row)
        ),
        reverse=True,
    )

    for candidate in candidates:
        if any(
            looks_like_decision(row)
            for row in candidate
        ):
            return candidate

    return []


def find_total(obj):
    total_keys = [
        "recordsTotal",
        "totalCount",
        "totalElements",
        "toplamKayit",
        "toplamKayıt",
        "toplam",
        "total",
        "count",
    ]

    def walk(value):
        if isinstance(value, dict):
            for key in total_keys:
                if key in value:
                    number = parse_number(value[key])

                    if number is not None:
                        return number

            for child in value.values():
                number = walk(child)

                if number is not None:
                    return number

        elif isinstance(value, list):
            for child in value:
                number = walk(child)

                if number is not None:
                    return number

        return None

    return walk(obj)


# ============================================================
# KARAR KAYDI NORMALİZASYONU
# ============================================================

def normalize_row(row):
    decision_id = pick(
        row,
        "id",
        "dokumanId",
        "dokumanID",
        "documentId",
        "documentID",
        "kararId",
        "kararID",
    )

    chamber = pick(
        row,
        "daire",
        "daireAdi",
        "daireAd",
        "birimAdi",
        "birim",
        "birimYrgKurulDaire",
        "kurulDaire",
        "mahkeme",
    )

    esas = pick(
        row,
        "esasNo",
        "esas",
        "esasNumarasi",
        "esasNumarası",
        "esasNoStr",
    )

    karar = pick(
        row,
        "kararNo",
        "karar",
        "kararNumarasi",
        "kararNumarası",
        "kararNoStr",
    )

    tarih = pick(
        row,
        "kararTarihi",
        "kararTarih",
        "kararTarihiStr",
        "tarih",
    )

    if not esas:
        esas_yil = pick(
            row,
            "esasYil",
            "esasYili",
            "esasYılı",
        )

        esas_sira = pick(
            row,
            "esasSiraNo",
            "esasSıraNo",
        )

        if esas_yil and esas_sira:
            esas = f"{esas_yil}/{esas_sira}"

    if not karar:
        karar_yil = pick(
            row,
            "kararYil",
            "kararYili",
            "kararYılı",
        )

        karar_sira = pick(
            row,
            "kararSiraNo",
            "kararSıraNo",
        )

        if karar_yil and karar_sira:
            karar = f"{karar_yil}/{karar_sira}"

    return {
        "id": clean(decision_id),
        "daire": clean(chamber),
        "esas": clean(esas),
        "karar": clean(karar),
        "tarih": clean(tarih),
    }


# ============================================================
# ARAMA PAYLOAD
# ============================================================

def build_search_data(
    query,
    chamber="",
    start_date="",
    end_date="",
    esas_year="",
    esas_first="",
    esas_last="",
    karar_year="",
    karar_first="",
    karar_last="",
):
    data = {
        "arananKelime": query,
        "birimYrgKurulDaire": chamber,
        "esasYil": esas_year,
        "esasIlkSiraNo": esas_first,
        "esasSonSiraNo": esas_last,
        "kararYil": karar_year,
        "kararIlkSiraNo": karar_first,
        "kararSonSiraNo": karar_last,
        "baslangicTarihi": start_date,
        "bitisTarihi": end_date,
        "siralama": "3",
        "siralamaDirection": "desc",
    }

    return {
        key: value
        for key, value in data.items()
        if value not in ("", None, "ALL")
    }


def payload_variants(data, page, page_size):
    return {
        1: {
            "data": data,
            "pageSize": page_size,
            "pageNumber": page,
        },

        2: {
            "data": {
                **data,
                "pageSize": page_size,
                "pageNumber": page,
            }
        },

        3: {
            **data,
            "pageSize": page_size,
            "pageNumber": page,
        },
    }


def valid_response(obj):
    rows = find_rows(obj)
    total = find_total(obj)

    if total is not None:
        return True

    if rows:
        return True

    text = json.dumps(
        obj,
        ensure_ascii=False,
    ).lower()

    return any(
        token in text
        for token in [
            "success",
            "adalet_success",
            "recordstotal",
            "başarılı",
            "basarili",
        ]
    )


# ============================================================
# RESMÎ KAYNAK ARAMA
# ============================================================

@st.cache_data(
    ttl=600,
    show_spinner=False,
)
def search_source(
    source,
    query,
    page,
    page_size,
    chamber,
    start_date,
    end_date,
    esas_year,
    esas_first,
    esas_last,
    karar_year,
    karar_first,
    karar_last,
):
    cfg = SOURCES[source]

    data = build_search_data(
        query,
        chamber,
        start_date,
        end_date,
        esas_year,
        esas_first,
        esas_last,
        karar_year,
        karar_first,
        karar_last,
    )

    variants = payload_variants(
        data,
        page,
        page_size,
    )

    state = runtime_state()
    known_schema = state["schema"].get(source)

    order = []

    if known_schema:
        order.append(known_schema)

    for schema in [1, 2, 3]:
        if schema not in order:
            order.append(schema)

    errors = []

    for schema in order:
        try:
            response = request_source(
                source,
                "POST",
                cfg["search"],
                json=variants[schema],
            )

            obj = response.json()

            if not valid_response(obj):
                errors.append(
                    f"Şema {schema}: "
                    "cevap geldi fakat karar yapısı tanınmadı."
                )
                continue

            state["schema"][source] = schema

            raw_rows = find_rows(obj)

            rows = [
                normalize_row(row)
                for row in raw_rows
                if isinstance(row, dict)
            ]

            total = find_total(obj)

            if total is None:
                total = len(rows)

            return {
                "ok": True,
                "rows": rows,
                "total": total,
                "schema": schema,
                "error": "",
            }

        except Exception as exc:
            errors.append(
                f"Şema {schema}: "
                f"{type(exc).__name__}: {exc}"
            )

    return {
        "ok": False,
        "rows": [],
        "total": 0,
        "schema": None,
        "error": "\n".join(errors),
    }


# ============================================================
# KARAR METNİ
# ============================================================

def extract_decision_text(raw):
    raw = html.unescape(raw or "")

    soup = BeautifulSoup(
        raw,
        "html.parser",
    )

    for tag in soup([
        "script",
        "style",
        "meta",
        "link",
        "noscript",
    ]):
        tag.decompose()

    candidates = []

    for tag in soup.find_all(True):
        text = tag.get_text(
            "\n",
            strip=True,
        )

        if len(text) >= 300:
            candidates.append(text)

    whole = soup.get_text(
        "\n",
        strip=True,
    )

    if whole:
        candidates.append(whole)

    decoded = html.unescape(whole)

    if "<" in decoded and ">" in decoded:
        soup2 = BeautifulSoup(
            decoded,
            "html.parser",
        )

        nested = soup2.get_text(
            "\n",
            strip=True,
        )

        if nested:
            candidates.append(nested)

    if not candidates:
        plain = re.sub(
            r"<br\s*/?>",
            "\n",
            raw,
            flags=re.I,
        )

        plain = re.sub(
            r"<[^>]+>",
            "",
            plain,
        )

        candidates.append(
            html.unescape(plain)
        )

    text = max(
        candidates,
        key=len,
    )

    junk = {
        "SUCCESS",
        "ADALET_SUCCESS",
        "İşlem başarıyla gerçekleştirildi!",
    }

    lines = []

    for line in text.splitlines():
        line = re.sub(
            r"[ \t]+",
            " ",
            line,
        ).strip()

        if not line:
            continue

        if line in junk:
            continue

        lines.append(line)

    return "\n".join(lines)


@st.cache_data(
    ttl=3600,
    show_spinner=False,
)
def get_decision_text(
    source,
    decision_id,
    query,
):
    if not decision_id:
        return ""

    cfg = SOURCES[source]

    url = cfg["doc"].format(
        id=quote(
            str(decision_id),
            safe="",
        ),
        term=quote(
            str(query or ""),
            safe="",
        ),
    )

    try:
        response = request_source(
            source,
            "GET",
            url,
        )

        return extract_decision_text(
            response.text
        )[:500000]

    except Exception:
        return ""


# ============================================================
# VURGULAMA
# ============================================================

def highlight_terms(query):
    query = clean(query)

    if not query:
        return []

    quoted = re.findall(
        r'"([^"]+)"',
        query,
    )

    if quoted:
        return [
            clean(x)
            for x in quoted
            if clean(x)
        ]

    return re.findall(
        r"[0-9A-Za-zÇĞİÖŞÜçğıöşü]+",
        query,
    )


def highlight_html(
    decision_text,
    query,
):
    safe_text = html.escape(
        decision_text
    )

    terms = highlight_terms(query)
    terms.sort(
        key=len,
        reverse=True,
    )

    for term in terms:
        escaped = html.escape(term)

        pattern = re.compile(
            re.escape(escaped),
            flags=re.I,
        )

        safe_text = pattern.sub(
            lambda match:
            (
                '<mark class="aranan">'
                + match.group(0)
                + "</mark>"
            ),
            safe_text,
        )

    safe_text = safe_text.replace(
        "\n",
        "<br>",
    )

    return f"""
    <div class="karar-metni">
        {safe_text}
    </div>
    """


# ============================================================
# WORD
# ============================================================

def make_word(
    source,
    row,
    decision_text,
    query,
):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    doc.add_heading(
        f"{source} Kararı",
        level=1,
    )

    doc.add_paragraph(
        f"Daire / Kurul: {row.get('daire') or '-'}"
    )

    doc.add_paragraph(
        f"Esas No: {row.get('esas') or '-'}"
    )

    doc.add_paragraph(
        f"Karar No: {row.get('karar') or '-'}"
    )

    doc.add_paragraph(
        f"Karar Tarihi: {row.get('tarih') or '-'}"
    )

    doc.add_paragraph(
        f"Arama: {query or '-'}"
    )

    doc.add_heading(
        "Karar Metni",
        level=2,
    )

    for paragraph in re.split(
        r"\n{2,}",
        decision_text,
    ):
        paragraph = paragraph.strip()

        if paragraph:
            doc.add_paragraph(
                paragraph
            )

    buffer = io.BytesIO()
    doc.save(buffer)

    return buffer.getvalue()


def word_name(source, row):
    source_name = (
        "Yargitay"
        if source == "Yargıtay"
        else "UYAP"
    )

    esas = (
        row.get("esas")
        or "Esas"
    ).replace("/", "_")

    karar = (
        row.get("karar")
        or "Karar"
    ).replace("/", "_")

    return (
        f"{source_name}_"
        f"E_{esas}_"
        f"K_{karar}.docx"
    )


# ============================================================
# TASARIM
# ============================================================

st.markdown(
    """
<style>

.block-container {
    max-width: 1750px;
    padding-top: 1.2rem;
}

.karar-metni {
    font-family: Georgia, "Times New Roman", serif;
    font-size: 17px;
    line-height: 1.72;
    height: 650px;
    overflow-y: auto;
    padding: 25px;
    border: 1px solid rgba(130,130,130,.3);
    border-radius: 12px;
    background: rgba(128,128,128,.035);
}

.aranan {
    background-color: #ffd900;
    color: #111;
    font-weight: 600;
    text-decoration: underline 2px #d18e00;
    text-underline-offset: 3px;
    padding: 0 2px;
}

div[data-testid="stMetric"] {
    border: 1px solid rgba(128,128,128,.2);
    border-radius: 12px;
    padding: 10px;
}

</style>
""",
    unsafe_allow_html=True,
)


# ============================================================
# ARAYÜZ
# ============================================================

st.title("⚖️ Yargı Kararı Bulucu")

st.caption(
    "Yargıtay ve UYAP Emsal kararlarında "
    "doğrudan arama ve karar inceleme."
)


search_left, search_right = st.columns(
    [7, 1]
)

with search_left:
    query = st.text_input(
        "Arama",
        placeholder=(
            'haksız tahrik   '
            'veya   '
            '"haksız tahrik"'
        ),
        label_visibility="collapsed",
    )

with search_right:
    search_button = st.button(
        "🔎 Ara",
        type="primary",
        use_container_width=True,
    )


with st.expander("🔍 Detaylı Arama"):
    col1, col2, col3 = st.columns(3)

    with col1:
        chamber = st.text_input(
            "Daire / Kurul",
            placeholder="Örn: 1. Ceza Dairesi",
        )

        esas_year = st.text_input(
            "Esas yılı",
            placeholder="2024",
        )

        esas_first = st.text_input(
            "Esas ilk sıra no",
        )

    with col2:
        start_date = st.text_input(
            "Başlangıç tarihi",
            placeholder="GG.AA.YYYY",
        )

        karar_year = st.text_input(
            "Karar yılı",
            placeholder="2025",
        )

        karar_first = st.text_input(
            "Karar ilk sıra no",
        )

    with col3:
        end_date = st.text_input(
            "Bitiş tarihi",
            placeholder="GG.AA.YYYY",
        )

        esas_last = st.text_input(
            "Esas son sıra no",
        )

        karar_last = st.text_input(
            "Karar son sıra no",
        )


# ============================================================
# SESSION STATE
# ============================================================

if "active_search" not in st.session_state:
    st.session_state.active_search = None

if "pages" not in st.session_state:
    st.session_state.pages = {
        "Yargıtay": 1,
        "UYAP Emsal": 1,
    }

if "selected" not in st.session_state:
    st.session_state.selected = {}


# ============================================================
# YENİ ARAMA
# ============================================================

if search_button:
    st.session_state.active_search = {
        "query": query,
        "chamber": chamber,
        "start_date": start_date,
        "end_date": end_date,
        "esas_year": esas_year,
        "esas_first": esas_first,
        "esas_last": esas_last,
        "karar_year": karar_year,
        "karar_first": karar_first,
        "karar_last": karar_last,
    }

    st.session_state.pages = {
        "Yargıtay": 1,
        "UYAP Emsal": 1,
    }

    st.session_state.selected = {}

    st.rerun()


active = st.session_state.active_search


# ============================================================
# ARAMA SONUÇLARI
# ============================================================

if active:
    if not any(
        clean(value)
        for value in active.values()
    ):
        st.warning(
            "Bir arama kriteri gir."
        )
        st.stop()

    PAGE_SIZE = 10

    # --------------------------------------------------------
    # KRİTİK DÜZELTME:
    # session_state worker thread içinde kullanılmıyor.
    # --------------------------------------------------------

    page_numbers = {
        "Yargıtay": int(
            st.session_state.pages.get(
                "Yargıtay",
                1,
            )
        ),
        "UYAP Emsal": int(
            st.session_state.pages.get(
                "UYAP Emsal",
                1,
            )
        ),
    }

    # active içeriğini de normal dict'e kopyala.
    active_data = dict(active)

    def search_job(
        source,
        page_number,
        active_copy,
    ):
        try:
            result = search_source(
                source,
                active_copy["query"],
                page_number,
                PAGE_SIZE,
                active_copy["chamber"],
                active_copy["start_date"],
                active_copy["end_date"],
                active_copy["esas_year"],
                active_copy["esas_first"],
                active_copy["esas_last"],
                active_copy["karar_year"],
                active_copy["karar_first"],
                active_copy["karar_last"],
            )

            return source, result

        except Exception as exc:
            return source, {
                "ok": False,
                "rows": [],
                "total": 0,
                "schema": None,
                "error": (
                    f"{type(exc).__name__}: "
                    f"{exc}"
                ),
            }

    results = {}

    with st.spinner(
        "Yargıtay ve UYAP Emsal taranıyor..."
    ):
        with ThreadPoolExecutor(
            max_workers=2
        ) as executor:

            futures = [
                executor.submit(
                    search_job,
                    source,
                    page_numbers[source],
                    active_data,
                )
                for source in [
                    "Yargıtay",
                    "UYAP Emsal",
                ]
            ]

            for future in as_completed(
                futures
            ):
                try:
                    source, result = (
                        future.result()
                    )

                    results[source] = result

                except Exception as exc:
                    # Worker'daki hata uygulamayı çökertmez.
                    results["Bilinmeyen"] = {
                        "ok": False,
                        "rows": [],
                        "total": 0,
                        "schema": None,
                        "error": (
                            f"{type(exc).__name__}: "
                            f"{exc}"
                        ),
                    }


    # ========================================================
    # KAYNAK EKRANI
    # ========================================================

    def render_source(
        source,
        result,
    ):
        if not result["ok"]:
            st.error(
                f"{source} sistemine "
                "bu oturumda erişilemedi."
            )

            with st.expander(
                "Teknik ayrıntı"
            ):
                st.code(
                    result["error"]
                    or
                    "Bilinmeyen bağlantı hatası"
                )

            return

        rows = result["rows"]
        total = result["total"]

        st.success(
            f"{total:,} adet karar bulundu."
            .replace(",", ".")
        )

        total_pages = max(
            1,
            (
                total
                + PAGE_SIZE
                - 1
            )
            // PAGE_SIZE
        )

        prev_col, page_col, next_col = (
            st.columns(
                [1, 2.2, 1]
            )
        )

        with prev_col:
            if st.button(
                "◀ Önceki",
                key=f"prev_{source}",
                disabled=(
                    st.session_state
                    .pages[source]
                    <= 1
                ),
                use_container_width=True,
            ):
                st.session_state.pages[
                    source
                ] -= 1

                st.session_state.selected.pop(
                    source,
                    None,
                )

                st.rerun()

        with page_col:
            selected_page = st.number_input(
                "Sayfa",
                min_value=1,
                max_value=total_pages,
                value=min(
                    st.session_state
                    .pages[source],
                    total_pages,
                ),
                step=1,
                key=f"page_{source}",
            )

            if (
                int(selected_page)
                !=
                st.session_state
                .pages[source]
            ):
                st.session_state.pages[
                    source
                ] = int(
                    selected_page
                )

                st.session_state.selected.pop(
                    source,
                    None,
                )

                st.rerun()

            st.caption(
                (
                    f"Sayfa "
                    f"{st.session_state.pages[source]:,}"
                    f" / "
                    f"{total_pages:,}"
                )
                .replace(",", ".")
            )

        with next_col:
            if st.button(
                "Sonraki ▶",
                key=f"next_{source}",
                disabled=(
                    st.session_state
                    .pages[source]
                    >= total_pages
                ),
                use_container_width=True,
            ):
                st.session_state.pages[
                    source
                ] += 1

                st.session_state.selected.pop(
                    source,
                    None,
                )

                st.rerun()

        if (
            rows
            and source
            not in st.session_state.selected
        ):
            first_with_id = next(
                (
                    row
                    for row in rows
                    if row.get("id")
                ),
                None,
            )

            if first_with_id:
                st.session_state.selected[
                    source
                ] = first_with_id

        left, right = st.columns(
            [0.45, 0.55],
            gap="large",
        )

        with left:
            st.markdown(
                "### Karar Listesi"
            )

            if not rows:
                st.info(
                    "Bu sayfada karar bulunamadı."
                )

            first_number = (
                (
                    st.session_state
                    .pages[source]
                    - 1
                )
                * PAGE_SIZE
            )

            for index, row in enumerate(
                rows,
                start=1,
            ):
                current = (
                    st.session_state
                    .selected.get(source)
                )

                is_selected = (
                    current
                    and current.get("id")
                    == row.get("id")
                )

                with st.container(
                    border=True
                ):
                    st.markdown(
                        f"**"
                        f"{first_number + index}. "
                        f"{row['daire'] or source}"
                        f"**"
                        f"  \n"
                        f"E. **{row['esas'] or '—'}**"
                        f" · "
                        f"K. **{row['karar'] or '—'}**"
                        f"  \n"
                        f"{row['tarih'] or '—'}"
                    )

                    if st.button(
                        (
                            "✓ Açık"
                            if is_selected
                            else
                            "📄 Önizle"
                        ),
                        key=(
                            f"{source}_"
                            f"{st.session_state.pages[source]}_"
                            f"{index}_"
                            f"{row.get('id')}"
                        ),
                        disabled=(
                            not bool(
                                row.get("id")
                            )
                        ),
                        use_container_width=True,
                    ):
                        st.session_state.selected[
                            source
                        ] = row

                        st.rerun()

        with right:
            st.markdown(
                "### Karar Önizleme"
            )

            selected_row = (
                st.session_state
                .selected
                .get(source)
            )

            if not selected_row:
                st.info(
                    "Soldaki listeden "
                    "bir karar seç."
                )
                return

            with st.spinner(
                "Karar metni yükleniyor..."
            ):
                decision_text = (
                    get_decision_text(
                        source,
                        selected_row["id"],
                        active_data["query"],
                    )
                )

            if not decision_text:
                st.warning(
                    "Karar listede bulundu ancak "
                    "tam metni bu istekte alınamadı."
                )
                return

            st.markdown(
                f"### "
                f"{selected_row['daire'] or source}"
            )

            st.markdown(
                f"**Esas:** "
                f"{selected_row['esas'] or '—'}"
                f" &nbsp;&nbsp; "
                f"**Karar:** "
                f"{selected_row['karar'] or '—'}"
                f" &nbsp;&nbsp; "
                f"**Tarih:** "
                f"{selected_row['tarih'] or '—'}"
            )

            st.markdown(
                highlight_html(
                    decision_text,
                    active_data["query"],
                ),
                unsafe_allow_html=True,
            )

            with st.expander(
                "📋 Kararı kopyala"
            ):
                st.caption(
                    "Kutunun içine tıkla → "
                    "Ctrl+A → Ctrl+C"
                )

                st.text_area(
                    "Tam karar metni",
                    value=decision_text,
                    height=450,
                    key=(
                        f"copy_"
                        f"{source}_"
                        f"{selected_row['id']}"
                    ),
                )

            word_bytes = make_word(
                source,
                selected_row,
                decision_text,
                active_data["query"],
            )

            download1, download2 = (
                st.columns(2)
            )

            with download1:
                st.download_button(
                    "⬇️ Word Olarak İndir",
                    data=word_bytes,
                    file_name=word_name(
                        source,
                        selected_row,
                    ),
                    mime=(
                        "application/"
                        "vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                    use_container_width=True,
                    key=(
                        f"word_"
                        f"{source}_"
                        f"{selected_row['id']}"
                    ),
                )

            with download2:
                st.download_button(
                    "⬇️ TXT Olarak İndir",
                    data=(
                        decision_text
                        .encode("utf-8")
                    ),
                    file_name="yargi_karari.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key=(
                        f"txt_"
                        f"{source}_"
                        f"{selected_row['id']}"
                    ),
                )


    # ========================================================
    # SEKME YAPISI
    # ========================================================

    tab_yargitay, tab_uyap, tab_all = (
        st.tabs([
            "⚖️ Yargıtay",
            "🏛️ UYAP Emsal",
            "📊 Birleşik Görünüm",
        ])
    )

    with tab_yargitay:
        render_source(
            "Yargıtay",
            results.get(
                "Yargıtay",
                {
                    "ok": False,
                    "rows": [],
                    "total": 0,
                    "error": "Sonuç alınamadı.",
                },
            ),
        )

    with tab_uyap:
        render_source(
            "UYAP Emsal",
            results.get(
                "UYAP Emsal",
                {
                    "ok": False,
                    "rows": [],
                    "total": 0,
                    "error": "Sonuç alınamadı.",
                },
            ),
        )

    with tab_all:
        st.markdown(
            "### Kaynak Özeti"
        )

        yargitay = results.get(
            "Yargıtay",
            {}
        )

        uyap = results.get(
            "UYAP Emsal",
            {}
        )

        col_y, col_u = st.columns(2)

        with col_y:
            if yargitay.get("ok"):
                st.metric(
                    "Yargıtay",
                    (
                        f"{yargitay.get('total', 0):,}"
                        .replace(",", ".")
                    ),
                )
            else:
                st.error(
                    "Yargıtay bağlantısı kurulamadı."
                )

        with col_u:
            if uyap.get("ok"):
                st.metric(
                    "UYAP Emsal",
                    (
                        f"{uyap.get('total', 0):,}"
                        .replace(",", ".")
                    ),
                )
            else:
                st.error(
                    "UYAP Emsal bağlantısı kurulamadı."
                )


# ============================================================
# BİLGİ
# ============================================================

with st.expander(
    "ℹ️ Arama nasıl çalışır?"
):
    st.markdown(
        """
### Tırnaksız arama

`haksız tahrik`

resmî sisteme tırnaksız gönderilir.

### Tırnaklı arama

`"haksız tahrik"`

tırnak işaretleri korunarak resmî sisteme gönderilir.

### Hız

Binlerce kararın tam metni aynı anda indirilmez.

Önce karar listesi alınır. Sadece seçtiğin kararın tam metni yüklenir.

### Vurgulama

`haksız tahrik` aramasında **haksız** ve **tahrik** ayrı ayrı vurgulanır.

`"haksız tahrik"` aramasında **haksız tahrik** ifadesi birlikte vurgulanır.
"""
    )